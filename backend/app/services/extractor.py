"""
Text extraction from uploaded PDF and DOCX files.
Includes magic bytes verification, corrupted/password-protected file handling,
scanned-image detection, and robust annotation offset mapping.
"""
from __future__ import annotations

import io
import re
import zipfile
from typing import Tuple


MAX_EXTRACT_CHARS = 12000  # Cap at ~10 pages of text for model safety


def validate_file_signature(file_bytes: bytes, filename: str) -> str:
    """
    Validate magic bytes / file signature to ensure the file contents
    match a legitimate PDF or DOCX file (prevents renamed .exe or .png attacks).
    Returns 'pdf' or 'docx', or raises ValueError with a friendly message.
    """
    if len(file_bytes) < 4:
        raise ValueError("The uploaded file is empty or corrupted. Please upload a valid resume.")

    # PDF magic byte check (%PDF)
    if file_bytes.startswith(b"%PDF"):
        return "pdf"

    # DOCX magic byte check (Zip container PK\x03\x04 or PK\x05\x06 or PK\x07\x08)
    if file_bytes.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
        # Verify it actually contains Word document structure
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                namelist = zf.namelist()
                if any("word/document.xml" in name or "[Content_Types].xml" in name for name in namelist):
                    return "docx"
        except Exception:
            raise ValueError(
                "This Word document appears corrupted or invalid. Try re-exporting it from Word or Google Docs."
            )
        return "docx"

    # If extension claims to be pdf/docx but signature fails:
    fn_lower = filename.lower()
    if fn_lower.endswith(".pdf"):
        raise ValueError(
            "This file has a .pdf extension, but does not appear to be a valid PDF. Make sure it wasn't renamed from another format."
        )
    elif fn_lower.endswith(".docx"):
        raise ValueError(
            "This file has a .docx extension, but does not appear to be a valid Word document. Make sure it wasn't renamed from another format."
        )
    else:
        raise ValueError("Only PDF and DOCX formats are accepted. Please upload a valid document.")


def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, bool]:
    """
    Extract text from PDF bytes using pdfplumber.
    Returns (extracted_text, was_truncated).
    Catches password-protected or unreadable PDFs cleanly.
    """
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        raise RuntimeError("pdfplumber is not installed on server.")

    text_parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            # Check for encrypted / password protected PDF
            if getattr(pdf, "is_encrypted", False):
                raise ValueError(
                    "This PDF is password-protected. Please remove the password protection and upload again."
                )

            if len(pdf.pages) == 0:
                raise ValueError("This PDF does not contain any pages. Please upload a valid resume.")

            for page_idx, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                except Exception:
                    continue

    except ValueError:
        raise
    except Exception as e:
        err_msg = str(e).lower()
        if "password" in err_msg or "encrypt" in err_msg:
            raise ValueError(
                "This PDF is password-protected. Please remove the password protection and upload again."
            )
        raise ValueError(
            "Couldn't read that PDF — it may be damaged or formatted unusually. Try exporting it fresh from Word or Docs."
        )

    full_text = "\n\n".join(text_parts).strip()

    # Detect scanned image PDF with no selectable text
    if len(full_text) < 100 or len(full_text.split()) < 20:
        raise ValueError(
            "Couldn't extract readable text from this PDF. Make sure your resume is an exported text PDF, not a scanned image or photo."
        )

    was_truncated = False
    if len(full_text) > MAX_EXTRACT_CHARS:
        full_text = full_text[:MAX_EXTRACT_CHARS].rsplit("\n", 1)[0]
        was_truncated = True

    return full_text, was_truncated


def extract_text_from_docx(file_bytes: bytes) -> tuple[str, bool]:
    """
    Extract text from DOCX bytes using python-docx.
    Returns (extracted_text, was_truncated).
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise RuntimeError("python-docx is not installed on server.")

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        raise ValueError(
            "Couldn't read that DOCX file. Try re-saving or exporting it as a PDF from Word or Google Docs."
        )

    lines: list[str] = []
    # Extract from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
        else:
            if lines and lines[-1] != "":
                lines.append("")

    # Also extract from tables if paragraphs were sparse
    for table in doc.tables:
        for row in table.rows:
            row_texts = list(dict.fromkeys(cell.text.strip() for cell in row.cells if cell.text.strip()))
            if row_texts:
                lines.append(" | ".join(row_texts))

    full_text = "\n".join(lines).strip()

    # Detect image-only / empty DOCX
    if len(full_text) < 100 or len(full_text.split()) < 20:
        raise ValueError(
            "This Word document doesn't contain enough readable text. Ensure your resume content is written as text rather than embedded in images."
        )

    was_truncated = False
    if len(full_text) > MAX_EXTRACT_CHARS:
        full_text = full_text[:MAX_EXTRACT_CHARS].rsplit("\n", 1)[0]
        was_truncated = True

    return full_text, was_truncated


def extract_text(filename: str, content_type: str, file_bytes: bytes) -> tuple[str, bool]:
    """
    Validate signature and dispatch to the correct extractor.
    Returns (extracted_text, was_truncated).
    """
    doc_type = validate_file_signature(file_bytes, filename)

    if doc_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif doc_type == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")


def map_quoted_text_to_offsets(
    quoted_texts: list[str], full_text: str
) -> list[tuple[int | None, int | None]]:
    """
    Find substring offsets for quoted_text values.
    Uses multi-stage matching: exact match -> case-insensitive -> normalized whitespace -> fuzzy fallback.
    Never crashes; returns (None, None) if no reliable match is found.
    """
    offsets: list[tuple[int | None, int | None]] = []
    lower_full = full_text.lower()
    # Normalized whitespace version
    norm_full = re.sub(r"\s+", " ", full_text).lower()

    for qt in quoted_texts:
        if not qt or not isinstance(qt, str):
            offsets.append((None, None))
            continue

        clean_qt = qt.strip().strip('"\'')
        if not clean_qt:
            offsets.append((None, None))
            continue

        # 1. Exact case-insensitive match
        idx = lower_full.find(clean_qt.lower())
        if idx != -1:
            offsets.append((idx, idx + len(clean_qt)))
            continue

        # 2. Match with normalized whitespace
        norm_qt = re.sub(r"\s+", " ", clean_qt).lower()
        norm_idx = norm_full.find(norm_qt)
        if norm_idx != -1:
            # Map back approximately to full_text index
            # Find the first word in full_text
            first_word = norm_qt.split()[0] if norm_qt.split() else ""
            if first_word:
                word_idx = lower_full.find(first_word)
                if word_idx != -1:
                    offsets.append((word_idx, min(len(full_text), word_idx + len(clean_qt))))
                    continue

        # 3. Substring match on first 30 chars
        if len(clean_qt) > 30:
            sub = clean_qt[:30].lower()
            sub_idx = lower_full.find(sub)
            if sub_idx != -1:
                offsets.append((sub_idx, min(len(full_text), sub_idx + len(clean_qt))))
                continue

        # Fallback: model paraphrased, drop annotation positioning safely
        offsets.append((None, None))

    return offsets
