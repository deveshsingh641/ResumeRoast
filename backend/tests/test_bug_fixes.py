"""
Regression tests for bug fixes:
1. DOCX table merged cell text deduplication
2. Database cursor None-safety in increment_usage and get_wall_entries
3. Roast router issues / strengths None-safety
"""
import io
import unittest
from unittest.mock import MagicMock, patch

from docx import Document
from fastapi.testclient import TestClient

from app.db import database
from app.main import app
from app.services import extractor


class TestBugFixes(unittest.TestCase):

    def test_merged_cells_docx_extractor(self):
        """Test that merged cells in DOCX tables do not duplicate text repeatedly."""
        doc = Document()
        doc.add_paragraph("Devesh Singh - Senior Software Engineer with 6 years of experience building distributed systems.")
        doc.add_paragraph("Core competencies include Python, FastAPI, React, PostgreSQL, Docker, Kubernetes and microservice architecture.")
        
        table = doc.add_table(rows=2, cols=3)
        # Merge cell 0 and cell 1 in row 0
        cell_0 = table.cell(0, 0)
        cell_1 = table.cell(0, 1)
        cell_0.merge(cell_1)
        cell_0.text = "Acme Corp Tech Lead (2021 - Present)"
        table.cell(0, 2).text = "Location: Bangalore"
        
        # Row 1 normal cells
        table.cell(1, 0).text = "Software Engineer"
        table.cell(1, 1).text = "Built high-performance payment gateways"
        table.cell(1, 2).text = "Handled 10M transactions daily"

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        extracted, _ = extractor.extract_text_from_docx(docx_bytes)
        
        # Check that 'Acme Corp Tech Lead' appears only once in that row line
        self.assertIn("Acme Corp Tech Lead (2021 - Present)", extracted)
        lines = [line.strip() for line in extracted.split("\n") if line.strip()]
        header_lines = [l for l in lines if "Acme Corp Tech Lead" in l]
        self.assertEqual(len(header_lines), 1)
        # Should NOT duplicate across merged cells
        self.assertNotIn("Acme Corp Tech Lead (2021 - Present) Acme Corp Tech Lead", header_lines[0])

    def test_increment_usage_handles_none_row(self):
        """Test that increment_usage safely falls back to 1 when cur.fetchone() returns None."""
        mock_conn = MagicMock()
        mock_cur = MagicMock()
        mock_conn.cursor.return_value.__enter__.return_value = mock_cur
        mock_cur.fetchone.return_value = None  # Simulating empty result

        with patch.object(database, "_get_conn") as mock_get_conn:
            mock_get_conn.return_value.__enter__.return_value = mock_conn
            count = database.increment_usage("192.168.1.100")
            self.assertEqual(count, 1)

    def test_get_wall_entries_handles_none_count_row(self):
        """Test that get_wall_entries safely handles None from count query."""
        mock_conn = MagicMock()
        mock_cur = MagicMock()
        mock_conn.cursor.return_value.__enter__.return_value = mock_cur
        
        # fetchall for entries
        mock_cur.fetchall.return_value = []
        # fetchone for count_query
        mock_cur.fetchone.return_value = None

        with patch.object(database, "_get_conn") as mock_get_conn:
            mock_get_conn.return_value.__enter__.return_value = mock_conn
            result = database.get_wall_entries(entry_type="shame", sort_by="recent", page=1, limit=10)
            self.assertEqual(result["total"], 0)
            self.assertEqual(result["items"], [])

    def test_get_roast_handles_null_issues_and_strengths(self):
        """Test that get_roast endpoint returns [] when issues/strengths are None in db."""
        fake_db_record = {
            "id": "11111111-1111-1111-1111-111111111111",
            "overall_score": 45,
            "band": "roasted",
            "one_line_verdict": "Total disaster",
            "issues": None,     # In DB, JSONB might be null
            "strengths": None,  # In DB, JSONB might be null
            "created_at": "2026-09-05T00:00:00Z",
            "candidate_name": "Test User",
            "unlocked": False,
        }

        with patch("app.routers.roast.database.get_roast", return_value=fake_db_record):
            client = TestClient(app)
            response = client.get("/api/roast/11111111-1111-1111-1111-111111111111")
            self.assertEqual(response.status_code, 200)
            data = response.json()
            self.assertEqual(data["issues"], [])
            self.assertEqual(data["strengths"], [])


if __name__ == "__main__":
    unittest.main()
